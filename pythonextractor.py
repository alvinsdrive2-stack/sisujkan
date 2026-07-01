"""
Extract soal from single docx file.
Usage: python extract_single.py <filepath> <type>
  type: ia04b | ia05 | ia05b | ia06
Output: JSON to stdout.
"""
import json, sys, re
from docx import Document

def extract_ia04b(docx_path):
    doc = Document(str(docx_path))
    soals = []
    for t in doc.tables:
        if len(t.rows) < 5 or len(t.columns) < 4:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        if "Aspek" not in c0 and "Lingkup" not in c0:
            continue
        for i in range(2, len(t.rows), 2):
            try:
                cells = [c.text.strip() for c in t.rows[i].cells]
            except:
                continue
            if not cells[0].strip() or not cells[0].strip()[0].isdigit():
                continue
            number = cells[0].strip().rstrip('.')
            lingkup = cells[1].strip() if len(cells) > 1 else ''
            soal_text = cells[2].strip() if len(cells) > 2 else ''
            soal_text = re.sub(r'^Soal:\s*', '', soal_text)
            kode_unit = cells[3].strip() if len(cells) > 3 else ''
            kode_unit = re.sub(r'^Kode Unit:\s*', '', kode_unit)
            soals.append({'no': int(number), 'soal': soal_text, 'lingkup': lingkup, 'kode_unit': kode_unit})
    return soals

def extract_ia05(docx_path):
    doc = Document(str(docx_path))
    soals = []
    for t in doc.tables:
        if len(t.rows) < 10 or len(t.columns) < 3:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        if "KUK" not in c0 and "SOAL" not in c0:
            continue
        r = 1
        while r < len(t.rows):
            try:
                cells = [c.text.strip() for c in t.rows[r].cells]
            except:
                r += 1; continue
            kuk = cells[0].strip()
            number = cells[1].strip()
            if not number.isdigit():
                r += 1; continue
            soal_text = cells[3].strip() if len(cells) > 3 and cells[3].strip() else (cells[2].strip() if len(cells) > 2 else '')
            options = {}
            for j in range(1, 5):
                if r + j < len(t.rows):
                    try:
                        oc = [c.text.strip() for c in t.rows[r + j].cells]
                        opt_letter = oc[2].strip().rstrip('.') if len(oc) > 2 else ''
                        opt_value = oc[3].strip() if len(oc) > 3 else ''
                        if opt_letter and opt_letter[0].lower() in 'abcd':
                            options[opt_letter[0].lower()] = opt_value
                    except:
                        pass
            soals.append({'no': int(number), 'soal': soal_text, 'kode_kuk': kuk, 'jawab_a': options.get('a', ''), 'jawab_b': options.get('b', ''), 'jawab_c': options.get('c', ''), 'jawab_d': options.get('d', '')})
            r += 5
    return soals

def extract_ia05b(docx_path):
    doc = Document(str(docx_path))
    answers = {}
    for t in doc.tables:
        if len(t.rows) < 5:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        if "Kunci Jawaban" not in c0 and "Jawaban" not in c0:
            continue
        for r in t.rows[2:]:
            try:
                cells = [c.text.strip() for c in r.cells]
            except:
                continue
            if cells[0].isdigit():
                answers[int(cells[0])] = cells[1].strip()
    return answers

def extract_ia06(docx_path):
    doc = Document(str(docx_path))
    soals = []
    for t in doc.tables:
        if len(t.rows) < 5 or len(t.columns) < 2:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        c1 = t.rows[0].cells[1].text.strip() if len(t.rows[0].cells) > 1 else ''
        if "KUK" not in c0 or ("ESAI" not in c0 and "ESAI" not in c1):
            continue
        for r in t.rows[1:]:
            try:
                cells = [c.text.strip() for c in r.cells]
            except:
                continue
            kuk = cells[0].strip()
            number = cells[1].strip() if len(cells) > 1 else ''
            soal_text = cells[2].strip() if len(cells) > 2 else ''
            if not number.isdigit(): continue
            soals.append({'no': int(number), 'soal': soal_text, 'kode_kuk': kuk})
    return soals

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(json.dumps({'error': 'Usage: extract_single.py <filepath> <type>'}), file=sys.stderr)
        sys.exit(1)
    filepath = sys.argv[1]
    dtype = sys.argv[2].lower()
    try:
        if dtype == 'ia04b':
            result = extract_ia04b(filepath)
        elif dtype == 'ia05':
            result = extract_ia05(filepath)
        elif dtype == 'ia05b':
            result = extract_ia05b(filepath)
        elif dtype == 'ia06':
            result = extract_ia06(filepath)
        else:
            result = {'error': f'Unknown type: {dtype}'}
        print(json.dumps(result, indent=2, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({'error': str(e)}), file=sys.stderr)
        sys.exit(1)
