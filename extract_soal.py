"""
Extract soal from docx bank soal.
Output: JSON per skema + merged JSON for all.
"""
import json, os, re
from pathlib import Path
from docx import Document
from tqdm import tqdm

ROOT = Path(r"C:\LSP\penyatuan\FE verif\gatensi-frontend\DOKUMENS\01. Bank Soal 150 Skema")
OUTPUT = Path(r"C:\LSP\penyatuan\FE verif\gatensi-frontend\extracted_soal")

def extract_ia04b(docx_path):
    """FR IA 04B - Lembar Periksa Kegiatan Terstruktur (Esai)"""
    doc = Document(str(docx_path))
    soals = []

    # Find main question table (largest table with questions)
    for t in doc.tables:
        if len(t.rows) < 5 or len(t.columns) < 4:
            continue
        # Check if first row has "Aspek Penilaian" header
        c0 = t.rows[0].cells[0].text.strip()
        if "Aspek" not in c0 and "Lingkup" not in c0:
            continue

        # Row 0-1 are headers, data starts at row 2
        # Questions come in pairs: (soal_row, detail_row)
        for i in range(2, len(t.rows), 2):
            try:
                cells = [c.text.strip() for c in t.rows[i].cells]
            except:
                continue
            if not cells[0].strip() or not cells[0].strip()[0].isdigit():
                continue

            number = cells[0].strip().rstrip('.')
            lingkup = cells[1].strip() if len(cells) > 1 else ''
            soal_text = cells[2].strip() if len(cells) > 2 else ''

            # Clean soal text
            soal_text = re.sub(r'^Soal:\s*', '', soal_text)

            # Extract kode unit from cell 3
            kode_unit = cells[3].strip() if len(cells) > 3 else ''
            kode_unit = re.sub(r'^Kode Unit:\s*', '', kode_unit)

            soals.append({
                'no': int(number),
                'soal': soal_text,
                'lingkup': lingkup,
                'kode_unit': kode_unit,
            })

    return soals

def extract_ia05(docx_path):
    """FR IA 05 - Pertanyaan Tertulis Pilihan Ganda"""
    doc = Document(str(docx_path))
    soals = []

    for t in doc.tables:
        if len(t.rows) < 10 or len(t.columns) < 3:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        if "KUK" not in c0 and "SOAL" not in c0:
            continue

        # Each question = 5 rows (KUK row + a/b/c/d rows)
        r = 1
        while r < len(t.rows):
            try:
                cells = [c.text.strip() for c in t.rows[r].cells]
            except:
                r += 1
                continue

            kuk = cells[0].strip()
            number = cells[1].strip()

            if not number.isdigit():
                r += 1
                continue

            # Question text (in cols 2 or 3)
            soal_text = cells[3].strip() if len(cells) > 3 and cells[3].strip() else (cells[2].strip() if len(cells) > 2 else '')

            # Next 4 rows = a, b, c, d
            options = {}
            for j in range(1, 5):
                if r + j < len(t.rows):
                    try:
                        oc = [c.text.strip() for c in t.rows[r + j].cells]
                        opt_letter = oc[2].strip().rstrip('.') if len(oc) > 2 else ''
                        opt_value = oc[3].strip() if len(oc) > 3 else ''
                        if opt_letter and opt_letter[0].lower() in 'abcd':
                            options[opt_letter[0].lower()] = opt_value
                    except:
                        pass

            soals.append({
                'no': int(number),
                'soal': soal_text,
                'kode_kuk': kuk,
                'jawab_a': options.get('a', ''),
                'jawab_b': options.get('b', ''),
                'jawab_c': options.get('c', ''),
                'jawab_d': options.get('d', ''),
            })

            r += 5

    return soals

def extract_ia05b(docx_path):
    """FR IA 05B - Kunci Jawaban PG"""
    doc = Document(str(docx_path))
    answers = {}

    for t in doc.tables:
        if len(t.rows) < 5:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        if "Kunci Jawaban" not in c0 and "Jawaban" not in c0:
            continue

        # Row 0-1 are headers, data from row 2
        for r in t.rows[2:]:
            try:
                cells = [c.text.strip() for c in r.cells]
            except:
                continue
            if cells[0].isdigit():
                answers[int(cells[0])] = cells[1].strip()

    return answers

def extract_ia06(docx_path):
    """FR IA 06 - Pertanyaan Tertulis Esai"""
    doc = Document(str(docx_path))
    soals = []

    for t in doc.tables:
        if len(t.rows) < 5 or len(t.columns) < 2:
            continue
        c0 = t.rows[0].cells[0].text.strip()
        c1 = t.rows[0].cells[1].text.strip() if len(t.rows[0].cells) > 1 else ''
        if "KUK" not in c0 or ("ESAI" not in c0 and "ESAI" not in c1):
            continue

        for r in t.rows[1:]:
            try:
                cells = [c.text.strip() for c in r.cells]
            except:
                continue
            kuk = cells[0].strip()
            number = cells[1].strip() if len(cells) > 1 else ''
            soal_text = cells[2].strip() if len(cells) > 2 else ''

            if not number.isdigit():
                continue

            soals.append({
                'no': int(number),
                'soal': soal_text,
                'kode_kuk': kuk,
            })

    return soals

def get_skema_name(folder):
    """Extract skema name from folder name"""
    name = folder.name
    # Remove leading number
    name = re.sub(r'^\d+\.\s*', '', name)
    return name.strip()

def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)

    skema_folders = sorted([f for f in ROOT.iterdir() if f.is_dir()])
    all_data = {}

    for folder in tqdm(skema_folders, desc="Processing skema"):
        skema_name = get_skema_name(folder)
        skema_data = {
            'nama': skema_name,
            'folder': folder.name,
            'ia04b': [],
            'ia05': [],
            'ia06': [],
        }

        # Find files
        files = list(folder.glob("*.docx"))
        ia04b_path = None
        ia05_path = None
        ia05b_path = None
        ia06_path = None

        for f in files:
            name = f.stem.upper()
            if 'IA 04B' in name or 'IA04B' in name:
                if 'KJ' not in name:
                    ia04b_path = f
            elif 'IA 05B' in name or 'IA05B' in name:
                ia05b_path = f
            elif 'IA 05' in name or 'IA05' in name:
                if 'B' not in name and 'C' not in name:
                    ia05_path = f
            elif 'IA 06' in name or 'IA06' in name:
                if 'B' not in name and 'C' not in name:
                    ia06_path = f

        # Extract
        if ia04b_path:
            skema_data['ia04b'] = extract_ia04b(ia04b_path)

        if ia05_path:
            skema_data['ia05'] = extract_ia05(ia05_path)
            if ia05b_path:
                answers = extract_ia05b(ia05b_path)
                for soal in skema_data['ia05']:
                    soal['jawaban'] = answers.get(soal['no'], '')

        if ia06_path:
            skema_data['ia06'] = extract_ia06(ia06_path)

        # Save per-skema JSON
        all_data[folder.name] = skema_data
        with open(OUTPUT / f"{folder.name}.json", 'w', encoding='utf-8') as f:
            json.dump(skema_data, f, indent=2, ensure_ascii=False)

    # Save all combined
    with open(OUTPUT / "ALL_SOAL.json", 'w', encoding='utf-8') as f:
        json.dump(all_data, f, indent=2, ensure_ascii=False)

    # Stats
    total_ia04b = sum(len(d['ia04b']) for d in all_data.values())
    total_ia05 = sum(len(d['ia05']) for d in all_data.values())
    total_ia06 = sum(len(d['ia06']) for d in all_data.values())

    print(f"\nDone! Extracted from {len(all_data)} skema:")
    print(f"  IA04B (esai): {total_ia04b} soals")
    print(f"  IA05 (PG): {total_ia05} soals")
    print(f"  IA06 (esai): {total_ia06} soals")
    print(f"\nOutput: {OUTPUT}/")

if __name__ == '__main__':
    main()
